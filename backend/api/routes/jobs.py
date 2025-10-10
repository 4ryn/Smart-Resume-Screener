from flask import Blueprint, request, jsonify
import os
import pandas as pd
from werkzeug.utils import secure_filename
import sys
import PyPDF2
from docx import Document
import re

# Import existing scripts
sys.path.append(os.path.join(os.path.dirname(__file__), '../../..'))
from load_jobs import load_job_descriptions
from jd_summarizer import process_job_descriptions
from backend.api.utils.db_helper import get_all_jobs

jobs_bp = Blueprint('jobs', __name__)

def clear_job_related_data():
    """Clear all job-related data for fresh upload"""
    import sqlite3
    from config import DB_PATH
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    print("🗑️  Clearing all job-related data...")
    cursor.execute("DELETE FROM jobs")
    cursor.execute("DELETE FROM shortlisted_candidates") 
    cursor.execute("UPDATE candidates SET match_score = NULL, matched_job_id = NULL")
    conn.commit()
    conn.close()
    print("✅ All job-related data cleared.")

UPLOAD_FOLDER = 'data'
ALLOWED_EXTENSIONS = {'csv', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            print(f"   📄 PDF has {len(pdf_reader.pages)} page(s)")
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    print(f"   ✓ Extracted page {page_num}: {len(page_text)} chars")
        return text.strip()
    except Exception as e:
        print(f"❌ Error extracting PDF text: {str(e)}")
        raise

def extract_text_from_docx(filepath):
    """Extract text from DOCX file"""
    try:
        doc = Document(filepath)
        print(f"   📄 DOCX has {len(doc.paragraphs)} paragraph(s)")
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        text = "\n".join(paragraphs)
        print(f"   ✓ Extracted {len(paragraphs)} non-empty paragraphs")
        return text.strip()
    except Exception as e:
        print(f"❌ Error extracting DOCX text: {str(e)}")
        raise

def smart_extract_job_info(text):
    """
    Intelligently extract job title and description from raw text.
    Handles multiple formats and patterns.
    """
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    if not lines:
        return "Untitled Position", text
    
    # Common job title indicators
    job_keywords = [
        'engineer', 'developer', 'manager', 'analyst', 'designer', 
        'architect', 'specialist', 'consultant', 'coordinator', 'lead',
        'director', 'administrator', 'scientist', 'researcher', 'officer',
        'executive', 'associate', 'assistant', 'technician', 'expert'
    ]
    
    job_title = None
    description_start_idx = 0
    
    # Strategy 1: Look for "Job Title:" or similar labels
    for i, line in enumerate(lines[:10]):
        line_lower = line.lower()
        if any(label in line_lower for label in ['job title:', 'position:', 'role:', 'title:']):
            # Extract the title (text after the label)
            job_title = re.sub(r'(job title|position|role|title)\s*:\s*', '', line, flags=re.IGNORECASE).strip()
            description_start_idx = i + 1
            break
    
    # Strategy 2: Look for lines with job keywords (first occurrence)
    if not job_title:
        for i, line in enumerate(lines[:10]):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in job_keywords):
                # Check if it's likely a title (short, no sentence structure)
                if len(line) < 100 and not line.endswith('.'):
                    job_title = line.strip(':-.,#*[](){}')
                    description_start_idx = i + 1
                    break
    
    # Strategy 3: First line is title if it's short enough
    if not job_title:
        first_line = lines[0].strip(':-.,#*[](){}')
        if len(first_line) < 100:
            job_title = first_line
            description_start_idx = 1
    
    # Fallback
    if not job_title:
        job_title = "Untitled Position"
        description_start_idx = 0
    
    # Extract description (everything after the title)
    description_lines = lines[description_start_idx:]
    
    # Remove common header patterns from description
    cleaned_description = []
    skip_patterns = ['job title', 'position', 'description:', 'responsibilities:', 'qualifications:']
    
    for line in description_lines:
        line_lower = line.lower()
        # Skip if it's just a label
        if any(pattern == line_lower.strip(':-.,#*') for pattern in skip_patterns):
            continue
        cleaned_description.append(line)
    
    job_description = '\n'.join(cleaned_description).strip()
    
    # If description is empty, use full text
    if not job_description:
        job_description = text
    
    return job_title, job_description

def process_csv_file(filepath):
    """Process CSV file and normalize to standard format"""
    try:
        print(f"📊 Processing CSV file...")
        df = pd.read_csv(filepath)
        
        print(f"   Columns found: {list(df.columns)}")
        print(f"   Rows: {len(df)}")
        
        # Check if CSV already has the correct format
        if 'Job Title' in df.columns and 'Job Description' in df.columns:
            print(f"   ✅ CSV already in correct format")
            
            # Ensure there's an empty third column
            if len(df.columns) == 2:
                df[''] = ''
            
            # Save normalized CSV
            output_path = os.path.join(UPLOAD_FOLDER, 'job_description.csv')
            df.to_csv(output_path, index=False)
            return len(df)
        
        # Try to find title and description columns with different names
        title_col = None
        desc_cols = []
        
        for col in df.columns:
            col_lower = col.lower().strip()
            if 'title' in col_lower or 'position' in col_lower or 'role' in col_lower:
                title_col = col
            elif any(keyword in col_lower for keyword in ['responsibilities', 'requirements', 'description', 'details', 'duties', 'qualifications']):
                desc_cols.append(col)
        
        # If we found a title column and multiple description columns, combine them
        if title_col and desc_cols:
            print(f"   ✓ Found title column: '{title_col}'")
            print(f"   ✓ Found description columns: {desc_cols}")
            
            # Combine all description columns into one
            job_descriptions = []
            for idx, row in df.iterrows():
                desc_parts = []
                for col in desc_cols:
                    val = str(row[col]).strip()
                    if val and val.lower() != 'nan':
                        desc_parts.append(f"{col}:\n{val}")
                
                combined_desc = "\n\n".join(desc_parts)
                job_descriptions.append(combined_desc)
            
            # Create normalized DataFrame
            normalized_df = pd.DataFrame({
                'Job Title': df[title_col],
                'Job Description': job_descriptions,
                '': ''
            })
            
            output_path = os.path.join(UPLOAD_FOLDER, 'job_description.csv')
            normalized_df.to_csv(output_path, index=False)
            print(f"   ✅ Combined {len(desc_cols)} columns into Job Description")
            return len(normalized_df)
        
        # If no clear columns found, treat first column as title, second as description
        if len(df.columns) >= 2:
            print(f"   ⚠️  Using first column as Title, second as Description")
            normalized_df = pd.DataFrame({
                'Job Title': df.iloc[:, 0],
                'Job Description': df.iloc[:, 1],
                '': ''
            })
            
            output_path = os.path.join(UPLOAD_FOLDER, 'job_description.csv')
            normalized_df.to_csv(output_path, index=False)
            return len(normalized_df)
        
        raise ValueError("CSV format not recognized. Expected columns: 'Job Title' and 'Job Description'")
        
    except Exception as e:
        print(f"❌ Error processing CSV: {str(e)}")
        raise

def convert_document_to_csv(text, original_filename):
    """Convert extracted text from PDF/DOCX to standardized CSV format"""
    try:
        print(f"🔄 Converting extracted text to CSV format...")
        
        # Smart extraction of job title and description
        job_title, job_description = smart_extract_job_info(text)
        
        print(f"   ✓ Extracted Job Title: {job_title[:60]}{'...' if len(job_title) > 60 else ''}")
        print(f"   ✓ Description length: {len(job_description)} characters")
        
        # Create standardized DataFrame
        df = pd.DataFrame({
            'Job Title': [job_title],
            'Job Description': [job_description],
            '': ['']
        })
        
        # Save to standard location
        output_path = os.path.join(UPLOAD_FOLDER, 'job_description.csv')
        df.to_csv(output_path, index=False)
        
        print(f"   ✅ Saved to: {output_path}")
        return 1  # One job added
        
    except Exception as e:
        print(f"❌ Error converting to CSV: {str(e)}")
        raise

def process_uploaded_file(filepath, filename):
    """
    Universal file processor - handles CSV, PDF, DOCX
    Converts everything to standardized CSV format: Job Title | Job Description | (empty)
    """
    file_ext = filename.rsplit('.', 1)[1].lower()
    
    print(f"\n{'='*60}")
    print(f"📁 Processing {file_ext.upper()} file: {filename}")
    print(f"{'='*60}\n")
    
    if file_ext == 'csv':
        # Process and normalize CSV
        job_count = process_csv_file(filepath)
        print(f"\n✅ CSV processing complete: {job_count} job(s) found\n")
        return job_count
    
    elif file_ext == 'pdf':
        # Extract text from PDF
        print(f"📄 Extracting text from PDF...")
        extracted_text = extract_text_from_pdf(filepath)
        
        if not extracted_text or len(extracted_text.strip()) < 50:
            raise ValueError(f"Could not extract sufficient text from PDF (only {len(extracted_text)} chars)")
        
        print(f"✅ Extracted {len(extracted_text)} characters from PDF\n")
        
        # Convert to CSV
        job_count = convert_document_to_csv(extracted_text, filename)
        return job_count
    
    elif file_ext == 'docx':
        # Extract text from DOCX
        print(f"📄 Extracting text from DOCX...")
        extracted_text = extract_text_from_docx(filepath)
        
        if not extracted_text or len(extracted_text.strip()) < 50:
            raise ValueError(f"Could not extract sufficient text from DOCX (only {len(extracted_text)} chars)")
        
        print(f"✅ Extracted {len(extracted_text)} characters from DOCX\n")
        
        # Convert to CSV
        job_count = convert_document_to_csv(extracted_text, filename)
        return job_count
    
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

@jobs_bp.route('/upload', methods=['POST'])
def upload_job_csv():
    """Upload and process job description (CSV, PDF, or DOCX)"""
    try:
        print(f"\n{'='*70}")
        print(f"📤 JOB UPLOAD REQUEST RECEIVED")
        print(f"{'='*70}")
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                'error': 'Invalid file type. Only CSV, PDF, and DOCX files are allowed'
            }), 400
        
        # Ensure upload directory exists
        os.makedirs(UPLOAD_FOLDER, exist_ok=True)
        
        # Get file extension
        file_ext = file.filename.rsplit('.', 1)[1].lower()
        filename = secure_filename(f'uploaded_job.{file_ext}')
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        
        print(f"📁 File: {file.filename}")
        print(f"💾 Saving to: {filepath}\n")
        file.save(filepath)
        
        # Clear existing job data
        clear_job_related_data()
        
        # Process file (convert to standardized CSV format)
        job_count = process_uploaded_file(filepath, file.filename)
        
        # Load jobs from the standardized CSV into database
        print(f"{'='*60}")
        print(f"📊 Loading jobs into database...")
        print(f"{'='*60}\n")
        load_job_descriptions()
        
        # Verify loaded jobs
        new_jobs = get_all_jobs()
        
        print(f"{'='*70}")
        print(f"✅ UPLOAD SUCCESSFUL")
        print(f"{'='*70}")
        print(f"File Type: {file_ext.upper()}")
        print(f"Jobs Loaded: {len(new_jobs)}")
        print(f"{'='*70}\n")
        
        return jsonify({
            'success': True,
            'message': f'Successfully processed {file_ext.upper()} file and loaded {len(new_jobs)} job(s)',
            'jobs_uploaded': len(new_jobs),
            'file_type': file_ext,
            'file_name': file.filename,
            'next_step': 'Call /api/jobs/summarize to process job descriptions with LLM'
        }), 200
        
    except Exception as e:
        print(f"\n{'='*60}")
        print(f"❌ UPLOAD ERROR")
        print(f"{'='*60}")
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        print(f"{'='*60}\n")
        return jsonify({'error': str(e)}), 500

@jobs_bp.route('/summarize', methods=['POST'])
def summarize_jobs():
    """Trigger JD summarization using LLM"""
    try:
        print(f"\n{'='*70}")
        print(f"🤖 STARTING JOB DESCRIPTION SUMMARIZATION")
        print(f"{'='*70}\n")
        
        process_job_descriptions()
        
        # Get updated jobs with summaries
        jobs = get_all_jobs()
        summarized_count = sum(1 for job in jobs if job.get('jd_summary'))
        
        print(f"\n{'='*70}")
        print(f"✅ SUMMARIZATION COMPLETE")
        print(f"{'='*70}")
        print(f"Total Jobs: {len(jobs)}")
        print(f"Summarized: {summarized_count}")
        print(f"{'='*70}\n")
        
        return jsonify({
            'success': True,
            'message': 'Job descriptions summarized successfully',
            'total_jobs': len(jobs),
            'summarized': summarized_count
        }), 200
        
    except Exception as e:
        print(f"\n❌ Summarization error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@jobs_bp.route('/', methods=['GET'])
def get_jobs():
    """Get all jobs"""
    try:
        jobs = get_all_jobs()
        response = jsonify({
            'data': jobs,
            'count': len(jobs),
            'timestamp': __import__('datetime').datetime.now().isoformat()
        })
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        return response, 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@jobs_bp.route('/clear', methods=['POST'])
def clear_all_data():
    """Clear all recruitment data for fresh start"""
    try:
        import sqlite3
        from config import DB_PATH
        
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        print("🗑️  Clearing ALL recruitment data...")
        cursor.execute("DELETE FROM jobs")
        cursor.execute("DELETE FROM candidates")
        cursor.execute("DELETE FROM shortlisted_candidates")
        conn.commit()
        conn.close()
        
        return jsonify({'message': 'All recruitment data cleared successfully'}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500